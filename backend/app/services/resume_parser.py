import io
import re
from typing import Optional

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes."""
    if not PDF_AVAILABLE:
        raise RuntimeError("pdfplumber not installed")
    
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")
    
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    
    return "\n".join(paragraphs)


def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    """Auto-detect file type and extract text."""
    ext = filename.rsplit('.', 1)[-1].lower()
    
    if ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext in ('docx', 'doc'):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_basic_info(text: str) -> dict:
    """Extract basic info from resume text using regex patterns."""
    info = {}
    
    # Email
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    info['email'] = emails[0] if emails else None
    
    # Phone
    phone_pattern = r'[\+]?[\d\s\-\(\)]{10,15}'
    phones = re.findall(phone_pattern, text)
    info['phone'] = phones[0].strip() if phones else None
    
    # LinkedIn
    linkedin_pattern = r'linkedin\.com/in/[\w\-]+'
    linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
    info['linkedin'] = linkedin[0] if linkedin else None
    
    # GitHub
    github_pattern = r'github\.com/[\w\-]+'
    github = re.findall(github_pattern, text, re.IGNORECASE)
    info['github'] = github[0] if github else None
    
    return info
